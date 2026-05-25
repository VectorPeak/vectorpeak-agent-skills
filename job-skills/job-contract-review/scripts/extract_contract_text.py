#!/usr/bin/env python3
"""Extract contract text from PDF, DOCX, TXT, MD, or RTF files."""

from __future__ import annotations

import argparse
import json
from pathlib import Path


TEXT_SUFFIXES = {".txt", ".md", ".rtf"}


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_pdf(path: Path) -> list[dict[str, str]]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise SystemExit("pdfplumber is required for PDF extraction: pip install pdfplumber") from exc

    chunks: list[dict[str, str]] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = normalize_text(text)
            if text:
                chunks.append({"source": f"page {index}", "text": text})

    if not chunks:
        raise SystemExit(
            "No extractable text found. This PDF is likely scanned/image-based and needs OCR first."
        )
    return chunks


def extract_docx(path: Path) -> list[dict[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("python-docx is required for DOCX extraction: pip install python-docx") from exc

    doc = Document(str(path))
    chunks: list[dict[str, str]] = []
    order = 1

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            chunks.append({"source": f"paragraph {order}", "text": text})
            order += 1

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            chunks.append({"source": f"table {table_index}", "text": "\n".join(rows)})

    if not chunks:
        raise SystemExit("No text found in DOCX file.")
    return chunks


def normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    cleaned_lines = [line for line in lines if line]
    return "\n".join(cleaned_lines).strip()


def extract_text(path: Path) -> dict[str, object]:
    if not path.exists():
        raise SystemExit(f"Input file does not exist: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        chunks = extract_pdf(path)
    elif suffix == ".docx":
        chunks = extract_docx(path)
    elif suffix in TEXT_SUFFIXES:
        text = normalize_text(read_text_file(path))
        if not text:
            raise SystemExit("No text found in text file.")
        chunks = [{"source": "text", "text": text}]
    else:
        raise SystemExit(f"Unsupported file type: {suffix}. Supported: .pdf, .docx, .txt, .md, .rtf")

    full_text = "\n\n".join(f"[{chunk['source']}]\n{chunk['text']}" for chunk in chunks)
    return {
        "source_file": str(path),
        "character_count": len(full_text),
        "chunks": chunks,
        "text": full_text,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Contract file path.")
    parser.add_argument("--output", help="Optional JSON output path. Prints JSON when omitted.")
    args = parser.parse_args()

    result = extract_text(Path(args.input))
    payload = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        output = Path(args.output)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(payload + "\n", encoding="utf-8")
    else:
        print(payload)


if __name__ == "__main__":
    main()
